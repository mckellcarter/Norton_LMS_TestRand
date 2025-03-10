import docx
import random

#XML tags
xml_paragraph = '<w:p>'
xml_paragraph_end = '</w:p>'
ques_tag = '<w:ilvl w:val="0"/>'
ans_tag = '<w:ilvl w:val="1"/>'
ans_key_tag = '<w:t>Answer:</w:t>'
word_text_tag = '<w:t>'
alt_word_text_tag = '<w:t xml:space="preserve">'
word_text_end_tag = '</w:t>'

class LMS_test_vers: 

  def __init__(self, docx_file): 
    self.questions = []
    self.__test_header__ = []
    self.__answer_key_header__ = []
    ##create a test object by reading in and parsing a docx generated from Norton's test maker for an LMS 
    
    #temporary structures
    q_temp = ''
    ans_temp = []
    q_temp_extra = []
    answer_key_count = -1  #assumes zero index to access questions
    in_key_answer = False  #use to flag when the next paragraph will have an answer
    
    d = docx.Document(docx_file) 
    ps = d.element.xml.split(xml_paragraph_end)
    for j,p in enumerate(ps): 
      ######process questions ###############
      if  ques_tag in p and not(self.__answer_key_header__):
        #if we have a non-indented list item and are outside the answer key, new question
        if q_temp == '':
        #if we haven't started a question
          q_temp = self.extract_text(p)
        else:
        #already have a question in process, time to push it and start a new one. 
          self.questions.append(LMS_question(q_temp, ans_temp, q_foot=q_temp_extra)) 
          q_temp = self.extract_text(p)
          ans_temp = []
          q_temp_extra = []
      elif ans_tag in p and q_temp != '': 
      #if we are in a question and have an indented list item
        ans_temp.append(self.extract_text(p))
      elif '<w:t>Answer Key</w:t>' in p:
      #begin Answer key, reset question gathering
        self.__answer_key_header__.append(self.extract_text(p))
        self.questions.append(LMS_question(q_temp, ans_temp, q_foot=q_temp_extra))
        q_temp = ''
        ans_temp = []    
        q_temp_extra = [] 
      elif q_temp and not(self.__answer_key_header__): 
        #processing a question but not an answer option so store to be dumped after answers
        q_temp_extra.append(self.extract_text(p))   
      elif self.__answer_key_header__ and ques_tag in p:
        #if in the answer key and detect an unindented list item 
        answer_key_count += 1 
        in_key_answer = True 
      elif in_key_answer and self.__answer_key_header__:
        if '<w:t>A</w:t>' in p: 
          key_answer = 0
        elif '<w:t>B</w:t>' in p: 
          key_answer = 1 
        elif '<w:t>C</w:t>' in p: 
          key_answer = 2
        elif '<w:t>D</w:t>' in p: 
          key_answer = 3 
        elif '<w:t>E</w:t>' in p: 
          key_answer = 4
        self.questions[answer_key_count].update_answer(key_answer)
        in_key_answer = False  #reset in_key_answer flag
      elif '<w:br w:type="page"/>' in p: 
         pass 
      elif not(self.__answer_key_header__):
        self.__test_header__.append(self.extract_text(p, concat=False))
      else: 
        self.__answer_key_header__.append(self.extract_text(p, concat=False))
        
  
  def add_paragraphs(self,old_list, new_item):
    if not(isinstance(old_list, list)): 
      old_list = [old_list]
    if isinstance(new_item, list): 
      for i in new_item: 
        old_list.append(i)
    elif not(isinstance(new_item, list)):
      old_list.append(new_item)
    return old_list 
      

  def print_new_test(self, fname, nQs=-1, shuffleQs=False, shuffleAs=False):
    ans_out = ['A', 'B', 'C', 'D', 'E']
    new_doc = docx.Document()
    key_doc = docx.Document()

    #add headers to new test and new answer key
    if isinstance(self.__test_header__, list): 
      for i in self.__test_header__:
        new_doc.add_paragraph(i)
    else: 
      new_doc.add_paragraph(self.__test_header__)
    if isinstance(self.__answer_key_header__, list):
      for i in self.__answer_key_header__:
        key_doc.add_paragraph(i)
    else: 
      key_doc.add_paragraph(self.__answer_key_header__)
    
    total_q_n = len(self.questions)
    new_inds = [i for i in range(total_q_n)]
    if shuffleQs: 
      random.shuffle(new_inds)
    if nQs > 0 and nQs < len(self.questions):
      new_inds = new_inds[:nQs]
    for i in new_inds: 
      tmp_question = self.questions[i]
      if shuffleAs: 
        tmp_question = self.questions[i].shuffle_answers()
      tmp_ans = tmp_question.print_question(new_doc, key_doc)
      key_doc.add_paragraph(ans_out[tmp_ans], style='List Number') 

    new_doc.save(fname)
    key_doc.save('key_'+fname)
 
  def extract_text(self, mystr, concat=True):
    #takes an xml paragraph as input. Extractts all text inside w:t tags
    ret_str = ''
    for j, q in enumerate(mystr.split(word_text_end_tag), concat):
      if word_text_tag in q or alt_word_text_tag in q:
        tmp_str = q.split(word_text_tag)[-1]
        tmp_str = tmp_str.split(alt_word_text_tag)[-1]
        if concat:
          if not(ret_str) and isinstance(ret_str, str):  
            ret_str += tmp_str
          elif not(tmp_str in ret_str[-1]) and isinstance(ret_str, str):
            ret_str += tmp_str
        else:
          if isinstance(ret_str, str):
            ret_str = []
          ret_str.append(tmp_str)
#    if not(concat) and isinstance(ret_str, list):
#      ret_str = '\n'.join(ret_str)
    return ret_str


class LMS_question: 
  def __init__(self,q, a, c=None, q_foot=[]): 
    self.question = q
    self.answers = a
    self.correct = c
    self.q_foot = q_foot

  def update_answer(self,c): 
    self.correct = c

  def shuffle_answers(self):
    #returns new question with shuffled answers 
    new_answers = self.answers.copy()
    corr_answer = self.answers[self.correct]
    random.shuffle(new_answers)
    new_correct = [i for i,v in enumerate(new_answers) if v==corr_answer][0] #find the index of the item matching the old correct answer (assumes unique answer entries) 
    return LMS_question(self.question, new_answers, c=new_correct, q_foot=self.q_foot)

  def print_question(self, testdoc, keydoc):
    testdoc.add_paragraph(self.question, style='List Number')
    for j,a in enumerate(self.answers): 
      ##method 1 
      #if j==0:
        #testdoc.add_paragraph(a, numbering: {
        #        reference: "padded-numbering-reference",
        #        level: 0,
        #        instance: 2}
      ##method 2 
      #testdoc.add_paragraph(a, style='List Number 2')
      ##method 3 
      if j == 0:
        ans_list = Paragraph_list(testdoc, a, True, 'Roman', level=2)
      else:
        ans_list.add_item(a, 2)
      

    if isinstance(self.q_foot, list): 
      for a in self.q_foot:
        testdoc.add_paragraph(a)
    else: 
      testdoc.add_paragraph(a)
    ret_ans = self.correct
    return ret_ans

class Paragraph_list(object):
  ##this solution to the list styles issue comes from rdt0086 on stack overflow
  #example use: 
  #mylist2 = Paragraph_List(document, 'Bullet Level 1 Item 1', False)
  #mylist2.add_item('Level 1 Item2',1)
  #mylist2.add_item('Level 2 Item1',2)
  #mylist2.add_item('Level 1 Item3',1)
  #mylist2.add_item('Level 2 Item1',2)
  #mylist2.add_item('Level 2 Item2',2)
  #mylist2.add_item('Level 2 Item3',2)
  #mylist2.add_item('Level 3 Item1',3)
  #mylist2.add_item('Level 4 Item1',4)
  #mylist2.add_item('Level 4 Item2',4)
  #mylist2.add_item('Level 2 Item4',2)

  #above use is:
  #ans_list = Paragraph_list(testdoc, a, True, 'ABC')
  #ans_list.add_item(ans2, 1)

  def __init__(self,doc,item1,ordered=False,style='',fmt=[[]],level=1): #args are: doc,item1,ordered,style,fmt,level (style, fmt, level optional)
    self.doc = doc
    self.item1 = item1
    self.ordered = ordered
    self.style = style
    self.fmt = fmt
    self.level = level
    self.place = {}
    self.place[level] = 0
 
    List_dict = {'Roman':[['I','II','III','IV','V','VI','VII','VIII','IIX','IX','X'],['A','B','C','D','E','F',
      'G','H','I','J','K','L'],['1','2','3','4','5','6','7','8','9','10'],['a','b','c','d','e','f','g','h',
      'i'],['i','ii','iii','iv','v','vi','vii','viii','iix','ix','x']],'ABC':[['A','B','C','D','E','F','G',
      'H','I','J','K','L'],['1','2','3','4','5','6','7','8','9','10'],['a','b','c','d','e','f','g','h','i'],
      ['i','ii','iii','iv','v','vi','vii','viii','iix','ix','x']],'123':[['1','2','3','4','5','6','7','8','9','10'
      ],['a','b','c','d','e','f','g','h','i'],['i','ii','iii','iv','v','vi','vii','viii','iix','ix','x']],'Bullet':[
      ['●','○','•','◦']]}
    if self.ordered == True:
      if not(self.style):
        self.fmt = List_dict['Roman']
      elif self.style == 'Custom':
        self.fmt = fmt
      else:
        self.fmt = List_dict[self.style]
        self.p = self.doc.add_paragraph("\t"*(self.level-1) + self.fmt[self.level-1][0]+'. ' + self.item1 + '\n')
    else:
      self.fmt = List_dict['Bullet']
      self.p = self.doc.add_paragraph("\t"*(self.level-1) +  self.fmt[self.level-1][0]+ ' ' + self.item1 + '\n')
      # self.doc.add_paragraph(item1, style=self.level)

  def add_item(self, item, level):
    self.level = level
    sp = "\t"
    sp = sp *(self.level-1)
    if self.level in self.place:
      self.place[self.level] += 1
    else:
      self.place[self.level] = 0
    if self.ordered ==True:
      self.p.add_run(sp + self.fmt[self.level - 1][self.place[self.level]] + '. ' + item + '\n')
    else:
      self.p.add_run(sp + self.fmt[0][self.level - 1]+ ' ' + item + '\n')
